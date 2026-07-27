"""Enterprise Report Generation Engine.

Generates professional reports with executive summaries, KPI tables,
insights, recommendations, methodology, and appendix.
Supports export to PDF, DOCX, HTML, and Markdown.
"""

from __future__ import annotations

import io
import json
import logging
import re
from datetime import datetime, timezone
from typing import Any

import pandas as pd
from sqlalchemy.orm import Session as DbSession

from ai.context_engine import EnterpriseAIContext, EnterpriseContextEngine
from ai.data_gatherer import DataGatherer
from ai.engines.executive_summary import ExecutiveSummaryEngine
from ai.engines.recommendation_engine import RecommendationEngine
from ai.gateway import AIGateway
from ai.models import AIReportGeneration
from ai.prompt_orchestrator import PromptOrchestrator, PromptTaskType

logger = logging.getLogger(__name__)


class EnterpriseReportEngine:
    """Generates professional reports from any dataset."""

    def __init__(self, db: DbSession | None = None):
        self.db = db
        self.gateway = AIGateway(db) if db else None
        self.context_engine = EnterpriseContextEngine(db)
        self.orchestrator = PromptOrchestrator()
        self.summary_engine = ExecutiveSummaryEngine(db)
        self.recommendation_engine = RecommendationEngine(db)

    def generate(
        self,
        report_type: str = "executive",
        title: str | None = None,
        df: pd.DataFrame | None = None,
        semantic_mappings: dict | None = None,
        industry: str = "unknown",
        user_id: int | None = None,
        context: EnterpriseAIContext | None = None,
        format: str = "markdown",
    ) -> dict:
        """Generate a professional report.

        Args:
            report_type: 'executive', 'monthly', 'annual', 'quality', 'performance'.
            title: Custom title. Auto-generated if None.
            df: DataFrame with the dataset.
            semantic_mappings: Semantic entity-to-column mappings.
            industry: Detected industry.
            user_id: User ID.
            context: Pre-built EnterpriseAIContext.
            format: Output format ('markdown', 'html', 'pdf', 'docx').

        Returns:
            Dict with id, report_type, title, content, summary, sections,
            format, created_at.
        """
        if context is None:
            context = self.context_engine.build(
                assistant_type="report_copilot",
                user_id=user_id,
                df=df,
                semantic_mappings=semantic_mappings,
                industry=industry,
            )

        if not title:
            title = self._default_title(report_type)

        gatherer = DataGatherer(df, context)
        report_data = gatherer.gather_for_report(report_type)

        exec_summary = self.summary_engine.generate(
            user_message=f"Generate an executive summary for the {report_type} report.",
            df=df,
            semantic_mappings=semantic_mappings,
            industry=industry,
            user_id=user_id,
            context=context,
            additional_data=report_data,
        )

        recommendations = self.recommendation_engine.generate(
            df=df,
            semantic_mappings=semantic_mappings,
            industry=industry,
            user_id=user_id,
            context=context,
            analysis_data=report_data,
        )

        sections = self._build_sections(
            report_type, title, report_data, exec_summary, recommendations, context
        )

        content = self._generate_content(title, sections, format)

        summary = exec_summary.get("executive_summary", "")
        methodology = self._generate_methodology(report_data, context)
        appendix = self._generate_appendix(report_data)

        content += f"\n\n## Methodology\n\n{methodology}"
        content += f"\n\n## Appendix\n\n{appendix}"

        report_id = None
        if self.db:
            try:
                report = AIReportGeneration(
                    report_type=report_type,
                    title=title,
                    content=content,
                    summary=summary,
                    sections=[s["title"] for s in sections],
                    format=format,
                    data_sources=report_data,
                    user_id=user_id,
                )
                self.db.add(report)
                self.db.commit()
                self.db.refresh(report)
                report_id = report.id
            except Exception as e:
                logger.warning(f"Failed to save report: {e}")

        exported = None
        if format in ("pdf", "docx", "html"):
            exported = self._export(content, format, title)

        return {
            "id": report_id,
            "report_type": report_type,
            "title": title,
            "content": content,
            "summary": summary,
            "sections": [s["title"] for s in sections],
            "methodology": methodology,
            "appendix": appendix,
            "format": format,
            "exported": exported,
            "executive_summary": exec_summary,
            "recommendations": recommendations,
            "created_at": str(datetime.now(timezone.utc).replace(tzinfo=None)),
        }

    def _default_title(self, report_type: str) -> str:
        """Generate a default title for the report."""
        now = datetime.now(timezone.utc)
        titles = {
            "executive": "Executive Summary Report",
            "monthly": f"Monthly Report - {now.strftime('%B %Y')}",
            "annual": f"Annual Report - {now.year}",
            "quality": "Data Quality Report",
            "performance": "Performance Report",
        }
        return titles.get(report_type, "Report")

    def _build_sections(
        self,
        report_type: str,
        title: str,
        data: dict,
        exec_summary: dict,
        recommendations: dict,
        context: EnterpriseAIContext,
    ) -> list[dict]:
        """Build the report sections."""
        sections = []

        # 1. Executive Summary
        sections.append({
            "title": "Executive Summary",
            "content": exec_summary.get("executive_summary", ""),
            "charts": [],
            "tables": [],
        })

        # 2. KPI Highlights
        kpi_highlights = exec_summary.get("kpi_highlights", [])
        if kpi_highlights:
            kpi_table = self._build_kpi_table(kpi_highlights)
            sections.append({
                "title": "KPI Highlights",
                "content": "The following key performance indicators were identified:",
                "charts": [],
                "tables": [kpi_table],
            })

        # 3. Main Drivers
        main_drivers = exec_summary.get("main_drivers", [])
        if main_drivers:
            sections.append({
                "title": "Main Drivers",
                "content": "\n".join(f"- {d}" for d in main_drivers),
                "charts": [],
                "tables": [],
            })

        # 4. Data Overview
        overall = data.get("overall", {})
        if overall:
            sections.append({
                "title": "Data Overview",
                "content": self._build_data_overview(overall),
                "charts": [],
                "tables": [],
            })

        # 5. Trend Analysis
        trends = data.get("time_trends", {})
        if trends:
            sections.append({
                "title": "Trend Analysis",
                "content": self._build_trend_section(trends),
                "charts": ["line"],
                "tables": [],
            })

        # 6. Top Contributors
        contributors = data.get("top_contributors", {})
        if contributors:
            sections.append({
                "title": "Top Contributors",
                "content": self._build_contributors_section(contributors),
                "charts": ["bar"],
                "tables": [],
            })

        # 7. Risks
        risks = exec_summary.get("risks", [])
        if risks:
            sections.append({
                "title": "Risks",
                "content": self._build_risks_section(risks),
                "charts": [],
                "tables": [],
            })

        # 8. Opportunities
        opportunities = exec_summary.get("opportunities", [])
        if opportunities:
            sections.append({
                "title": "Opportunities",
                "content": "\n".join(f"- {o}" for o in opportunities),
                "charts": [],
                "tables": [],
            })

        # 9. Recommendations
        recs = recommendations.get("recommendations", [])
        if recs:
            sections.append({
                "title": "Recommendations",
                "content": self._build_recommendations_section(recs),
                "charts": [],
                "tables": [],
            })

        return sections

    def _build_kpi_table(self, kpi_highlights: list[dict]) -> str:
        """Build a markdown table for KPI highlights."""
        header = "| Metric | Value | Change | Direction |"
        separator = "|--------|-------|--------|-----------|"
        rows = []
        for kpi in kpi_highlights:
            metric = kpi.get("metric", "")
            value = kpi.get("value", "")
            change = kpi.get("change", "")
            direction = kpi.get("direction", "")
            rows.append(f"| {metric} | {value} | {change} | {direction} |")
        return "\n".join([header, separator] + rows)

    def _build_data_overview(self, overall: dict) -> str:
        """Build the data overview section."""
        lines = [f"- **Total Records**: {overall.get('row_count', 0):,}"]
        lines.append(f"- **Total Columns**: {overall.get('column_count', 0)}")
        for key, val in overall.items():
            if key.startswith("total_") and isinstance(val, (int, float)):
                label = key.replace("total_", "Total ").title()
                lines.append(f"- **{label}**: {val:,.2f}")
            elif key.startswith("avg_") and isinstance(val, (int, float)):
                label = key.replace("avg_", "Average ").title()
                lines.append(f"- **{label}**: {val:,.2f}")
            elif key.startswith("unique_") and isinstance(val, int):
                label = key.replace("unique_", "Unique ").title()
                lines.append(f"- **{label}**: {val}")
        return "\n".join(lines)

    def _build_trend_section(self, trends: dict) -> str:
        """Build the trend analysis section."""
        lines = []
        for metric, values in trends.items():
            if values:
                first = values[0]["value"]
                last = values[-1]["value"]
                change = last - first
                pct = (change / first * 100) if first != 0 else 0
                direction = "increasing" if change > 0 else "decreasing" if change < 0 else "stable"
                lines.append(f"- **{metric.title()}**: {direction} ({pct:+.1f}%) from {first:,.2f} to {last:,.2f}")
        return "\n".join(lines) if lines else "No trend data available."

    def _build_contributors_section(self, contributors: dict) -> str:
        """Build the top contributors section."""
        lines = []
        for key, items in contributors.items():
            label = key.replace("_", " ").title()
            lines.append(f"\n**{label}**:")
            for item in items[:5]:
                dim_name = list(item.keys())[0]
                val_name = list(item.keys())[1]
                share = item.get("share", 0)
                lines.append(f"  - {item[dim_name]}: {item[val_name]:,.2f} ({share:.1f}%)")
        return "\n".join(lines) if lines else "No contributor data available."

    def _build_risks_section(self, risks: list) -> str:
        """Build the risks section."""
        lines = []
        for risk in risks:
            if isinstance(risk, dict):
                desc = risk.get("risk", "")
                severity = risk.get("severity", "")
                evidence = risk.get("evidence", "")
                lines.append(f"- **{desc}** (Severity: {severity})")
                if evidence:
                    lines.append(f"  - Evidence: {evidence}")
            else:
                lines.append(f"- {risk}")
        return "\n".join(lines) if lines else "No risks identified."

    def _build_recommendations_section(self, recs: list[dict]) -> str:
        """Build the recommendations section."""
        lines = []
        for rec in recs:
            action = rec.get("action", "")
            priority = rec.get("priority", "")
            impact = rec.get("expected_impact", "")
            feasibility = rec.get("feasibility", "")
            lines.append(f"- **{action}**")
            lines.append(f"  - Priority: {priority}")
            lines.append(f"  - Expected Impact: {impact}")
            lines.append(f"  - Feasibility: {feasibility}")
        return "\n".join(lines) if lines else "No recommendations available."

    def _generate_content(self, title: str, sections: list[dict], format: str) -> str:
        """Generate the full report content in the requested format."""
        lines = [f"# {title}\n"]
        for section in sections:
            lines.append(f"\n## {section['title']}\n")
            if section.get("content"):
                lines.append(section["content"])
            if section.get("tables"):
                for table in section["tables"]:
                    lines.append(f"\n{table}\n")
            if section.get("charts"):
                chart_types = ", ".join(section["charts"])
                lines.append(f"\n*[Chart: {chart_types}]*\n")
        return "\n".join(lines)

    def _generate_methodology(self, data: dict, context: EnterpriseAIContext) -> str:
        """Generate the methodology section."""
        lines = ["This report was generated using the following methodology:\n"]
        lines.append(f"- **Data Source**: {context.dataset.name or 'Platform dataset'}")
        lines.append(f"- **Records Analyzed**: {data.get('overall', {}).get('row_count', 0):,}")
        lines.append(f"- **Columns Analyzed**: {data.get('overall', {}).get('column_count', 0)}")
        if context.industry.industry != "unknown":
            lines.append(f"- **Industry Context**: {context.industry.display_name or context.industry.industry}")
        lines.append("- **Analysis Methods**: Statistical aggregation, trend analysis, contribution analysis")
        lines.append("- **AI Model**: Enterprise AI Decision Support System")
        lines.append("- **Report Generated**: " + str(datetime.now(timezone.utc).replace(tzinfo=None)))
        if data.get("data_sources"):
            lines.append(f"- **Data Sources**: {len(data['data_sources'])} source(s)")
        return "\n".join(lines)

    def _generate_appendix(self, data: dict) -> str:
        """Generate the appendix with detailed data tables."""
        lines = ["### Detailed Statistics\n"]
        stats = data.get("numeric_stats", {})
        if stats:
            for col, col_stats in stats.items():
                lines.append(f"\n**{col}**:")
                for stat_name, stat_val in col_stats.items():
                    lines.append(f"  - {stat_name.title()}: {stat_val:,.2f}")
        overall = data.get("overall", {})
        if overall:
            lines.append("\n### Overall Summary\n")
            for key, val in overall.items():
                if isinstance(val, (int, float)):
                    lines.append(f"  - {key.replace('_', ' ').title()}: {val:,.2f}")
        return "\n".join(lines)

    def _export(self, content: str, format: str, title: str) -> dict | None:
        """Export the report to the requested format."""
        try:
            if format == "html":
                return {"format": "html", "content": self._to_html(content, title)}
            elif format == "pdf":
                return {"format": "pdf", "content": self._to_pdf(content, title)}
            elif format == "docx":
                return {"format": "docx", "content": self._to_docx(content, title)}
        except Exception as e:
            logger.warning(f"Export to {format} failed: {e}")
            return {"format": format, "error": str(e)}
        return None

    def _to_html(self, content: str, title: str) -> str:
        """Convert markdown content to HTML."""
        try:
            import markdown as md
            html_body = md.markdown(content, extensions=["tables", "fenced_code"])
            return f"<!DOCTYPE html>\n<html>\n<head>\n<title>{title}</title>\n<style>\nbody {{ font-family: Arial, sans-serif; margin: 40px; }}\ntable {{ border-collapse: collapse; width: 100%; }}\nth, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}\nth {{ background-color: #f2f2f2; }}\n</style>\n</head>\n<body>\n{html_body}\n</body>\n</html>"
        except ImportError:
            return f"<pre>{content}</pre>"

    def _to_pdf(self, content: str, title: str) -> bytes:
        """Convert markdown content to PDF."""
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)

        # Strip markdown formatting for PDF
        clean = re.sub(r'[#*|_`>]', '', content)
        clean = clean.replace('\n\n', '\n')

        pdf.multi_cell(0, 6, clean[:5000])  # Limit content for PDF
        return pdf.output(dest="S")

    def _to_docx(self, content: str, title: str) -> bytes:
        """Convert markdown content to DOCX."""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(title, 0)

            for line in content.split("\n"):
                clean = re.sub(r'[#*|_`>]', '', line).strip()
                if clean:
                    if line.startswith("# "):
                        doc.add_heading(clean, 1)
                    elif line.startswith("## "):
                        doc.add_heading(clean, 2)
                    elif line.startswith("### "):
                        doc.add_heading(clean, 3)
                    elif line.startswith("- "):
                        doc.add_paragraph(clean, style="List Bullet")
                    else:
                        doc.add_paragraph(clean)

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
        except ImportError:
            return content.encode("utf-8")
